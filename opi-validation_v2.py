#!/usr/bin/python

from docx import Document
import time
import datetime
import os, sys
from pathlib import Path
import xml.etree.ElementTree as ET


CONREP_PÂTH='/tmp/conrep_current.dat'
FILE_PATH='/tmp/validation'+str(datetime.date.today())+'.docx'
URL='http://confluence.tid.es/page/viewpage.action?pageId=5275722'



#VALIDACION DE LA DOCUMENACION
validation_DOC = (	('Crear entrada RDM',('echo','0')),
					('Validar la ubicacion',('echo','0')),
					('Validar Nombre',('echo','0')),
					('Comprobar tabla de ips', ('echo','0')))
					
#VALIDACION DE LA ILO/IBMC
validation_ILOiBMC = (	('Verificar accesos WEB ILO/iBMC', ('echo','0')),
						('Verificar licencia ILO/iBMC',('echo','0')),
						('Crear usuarios: gvucdb/operfo',('echo','0')),
						('listar Usuarios ILO/iBMC', ('echo','0')))
#VALIDACION DE HW y SW						
validation_LINUX  = ( 	('contenido force',('sudo cat /opt/p2pcdn/conf-management/hieradata/force.yaml', '0')),
						('Errores Journal',('sudo journalctl -n 500','0')),
						('Rpm instalados',('sudo cdn-info.sh | grep tid-cdn','0')),
						('Estado Filesystem',('df -h --output=source,fstype,size,used,avail,pcent,target -x tmpfs -x devtmpfs','0')),
						('Modelo de HW',('sudo facter hw_model','0')
						('Estado de los discos',('sudo /opt/MegaRAID/storcli/storcli64 /c0 show','sudo /opt/smartstorageadmin/ssacli/bin/ssacli ctrl slot=0 ld all show detail')),
						('Revisar estado de la red',('sudo ip a','0')),
						('Revision zabbix',('sudo systemctl status zabbix-agent | head -4','0')),
						('Revision td-agent',('sudo systemctl status td-agent | head -4','0')),
						('Revision chronyd',('sudo systemctl status chronyd | head -4','0')),
						('Revision Procesos TID',('if [ $(sudo facter role) = "webcache" ] ; then sudo systemctl status tid-cdn-nginx |head -4; elif [ $(sudo facter role) = "tracker" ] ; then sudo systemctl status tid-cdn-pdns |head -4; sudo systemctl status tid-cdn-tracker-server | head -4 ; sudo systemctl status tid-cdn-tracker-worker | head -4; sudo systemctl status mysqld | head -4; else  sudo systemctl status tid-cdn-node | head -4; sudo systemctl status tid-cdn-ml | head -4; fi','0')),
						('Revision limites saturacion',('if [ $(sudo facter role) = "webcache" ] ; then sudo cat /opt/p2pcdn/etc/webcache_ctrl.conf | grep limit | grep outbound-bandwidth-limit ; elif [ $(sudo facter role) = "tracker" ] ; then echo "No Aplica" ; else  sudo cat /opt/p2pcdn/etc/node/nodeconfig.cfg ; fi','0')),
						('Revision sources Chrony',('sudo chronyc sources','0'),
						('Flujos en error', ('sudo cdn-check-flows.sh','0')),
						('Flujos en error', ('sudo op-powersaving.py','0')),
						('Expiracion pass TCDN ROOT',('sudo chage -l tcdnroot','0')))
						
def validation(self,secction,steps,document):
		document.add_paragraph('\n1. '+secction)
		table = document.add_table(rows=1, cols=4)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'ELEMENTO'
		hdr_cells[1].text = 'DESCRIPCION PRUEBA'
		hdr_cells[2].text = 'RESULTADO'
		hdr_cells[3].text = 'OBSERVACIONES'
		for prueba, cmd, resultado in steps:
			row_cells = table.add_row().cells
			row_cells[0].text = str(prueba)
			if cmd[1] == '0':
				row_cells[1].text = os.popen(cmd[0]).read()
			else:
				if 'Gen' in hw_model:
					row_cells[1].text = os.popen(cmd[0]).read()
				else:
					row_cells[1].text = os.popen(cmd[1]).read()
			row_cells[2].text = 'OK'
			row_cells[3].text = 'N/A'
		return document	

def main(self):
	path=FILE_PATH
	url=URL
	if Path(path):
		rm='sudo rm -rf '+path
		os.system(rm)

	document = Document()
	
	hostname = os.popen('hostname').read()
	typeElement = os.popen('sudo facter role').read()
	hw_model = os.popen('sudo facter hw_model').read()
	document.add_heading('Checklist elementos de Nueva Huella', 0)
	n = document.add_paragraph('Némonico; '+ str(hostname)+'\nTipo Elemento: '+ str(typeElement))
	u = document.add_paragraph(str(url))
	document = validation('VALIDACION DE LA DOCUMENACION',validation_DOC,document)
	document = validation('VALIDACION DE LA ILO/IBMC',validation_ILOiBMC,document)
	document = validation('VALIDACION DE HW y SW',validation_LINUX,document)
	
	document.save(path)
	
main()




	